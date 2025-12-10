"""
文件处理服务
支持PDF、DOCX、EPUB、Markdown等格式的内容提取和转换
"""
import os
import tempfile
import re
import logging
from pathlib import Path
from typing import Dict, Optional, Tuple
import sys

logger = logging.getLogger('common')

# 添加apps路径到sys.path
apps_path = Path(__file__).parent.parent.parent.parent / 'apps'
if str(apps_path) not in sys.path:
    sys.path.insert(0, str(apps_path))

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    HAS_EPUB = True
except ImportError:
    HAS_EPUB = False


def get_file_content(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    根据文件类型提取内容并转换为Markdown
    :param file_path: 文件路径
    :param file_name: 文件名
    :return: (文件内容, 转换后的文件名)
    """
    file_ext = Path(file_name).suffix.lower()
    logger.info(f'开始提取文件内容: {file_name}, 类型: {file_ext}')
    
    try:
        if file_ext == '.pdf':
            content, md_name = process_pdf(file_path, file_name)
        elif file_ext == '.docx':
            content, md_name = process_docx(file_path, file_name)
        elif file_ext == '.epub':
            content, md_name = process_epub(file_path, file_name)
        elif file_ext in ['.md', '.markdown', '.txt']:
            content, md_name = process_text(file_path, file_name)
        else:
            raise ValueError(f'不支持的文件格式: {file_ext}')
        
        logger.info(f'文件内容提取完成: {file_name}, 内容长度: {len(content)} 字符, Markdown文件名: {md_name}')
        return content, md_name
    except Exception as e:
        logger.error(f'提取文件内容失败: {file_name}, 错误: {str(e)}', exc_info=True)
        raise


def process_pdf(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    处理PDF文件，提取文本内容
    :param file_path: PDF文件路径
    :param file_name: 文件名
    :return: (Markdown内容, 转换后的文件名)
    """
    if not HAS_PYMUPDF:
        logger.error(f'PyMuPDF (fitz) 未安装，无法处理PDF文件: {file_name}')
        raise ImportError('PyMuPDF (fitz) 未安装，请运行: pip install PyMuPDF')
    
    try:
        logger.debug(f'开始处理PDF文件: {file_path}')
        doc = fitz.open(file_path)
        content = ""
        
        # 尝试从目录提取内容
        toc = doc.get_toc()
        logger.debug(f'PDF文件页数: {len(doc)}, 目录条目数: {len(toc) if toc else 0}')
        
        if toc and len(toc) > 0:
            # 按章节提取
            chapters = []
            for i, entry in enumerate(toc):
                level, title, start_page = entry
                start_page -= 1  # PyMuPDF页码从0开始
                
                # 确定结束页码
                if i + 1 < len(toc):
                    end_page = toc[i + 1][2] - 1
                else:
                    end_page = doc.page_count - 1
                
                # 提取章节内容
                chapter_content = ""
                for page_num in range(start_page, min(end_page + 1, doc.page_count)):
                    page = doc.load_page(page_num)
                    chapter_content += page.get_text()
                
                # 清理标题
                clean_title = re.sub(r'[^\w\s-]', '', title).strip()
                
                if chapter_content.strip():
                    # 根据级别添加Markdown标题
                    heading = '#' * min(level, 6)
                    chapters.append(f"{heading} {clean_title}\n\n{chapter_content.strip()}\n\n")
            
            content = "\n".join(chapters)
            logger.debug(f'从目录提取内容完成, 章节数: {len(chapters)}')
        else:
            # 没有目录，逐页提取
            logger.debug('PDF文件无目录，开始逐页提取')
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text and text.strip():
                    content += text + "\n\n"
        
        doc.close()
        
        # 清理内容
        content = content.replace('\0', '')  # 移除空字符
        content = re.sub(r'\n{3,}', '\n\n', content)  # 清理多余空行
        
        # 生成Markdown文件名
        md_file_name = file_name.replace('.pdf', '.md')
        
        logger.info(f'PDF处理完成: {file_name}, 提取内容长度: {len(content)} 字符')
        return content.strip(), md_file_name
    except Exception as e:
        logger.error(f'PDF处理失败: {file_name}, 错误: {str(e)}', exc_info=True)
        raise Exception(f'PDF处理失败: {str(e)}')


def process_docx(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    处理DOCX文件，转换为Markdown
    :param file_path: DOCX文件路径
    :param file_name: 文件名
    :return: (Markdown内容, 转换后的文件名)
    """
    if not HAS_DOCX:
        logger.error(f'python-docx 未安装，无法处理DOCX文件: {file_name}')
        raise ImportError('python-docx 未安装，请运行: pip install python-docx')
    
    try:
        logger.debug(f'开始处理DOCX文件: {file_path}')
        doc = Document(file_path)
        content_parts = []
        
        # 处理段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检查样式
            style_name = paragraph.style.name if paragraph.style else ''
            
            # 标题样式转换为Markdown标题
            if 'Heading' in style_name or '标题' in style_name:
                level = 1
                if '1' in style_name or '一' in style_name:
                    level = 1
                elif '2' in style_name or '二' in style_name:
                    level = 2
                elif '3' in style_name or '三' in style_name:
                    level = 3
                else:
                    level = min(int(re.search(r'\d+', style_name).group()) if re.search(r'\d+', style_name) else 1, 6)
                
                content_parts.append(f"{'#' * level} {text}\n")
            else:
                content_parts.append(f"{text}\n")
        
        # 处理表格
        for table in doc.tables:
            table_md = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_md.append("| " + " | ".join(cells) + " |")
            
            if table_md:
                # 添加表头分隔符
                if len(table_md) > 0:
                    header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                    table_md.insert(1, header_sep)
                
                content_parts.append("\n" + "\n".join(table_md) + "\n")
        
        content = "\n".join(content_parts)
        
        # 清理内容
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # 生成Markdown文件名
        md_file_name = file_name.replace('.docx', '.md')
        
        logger.info(f'DOCX处理完成: {file_name}, 提取内容长度: {len(content)} 字符')
        return content.strip(), md_file_name
    except Exception as e:
        logger.error(f'DOCX处理失败: {file_name}, 错误: {str(e)}', exc_info=True)
        raise Exception(f'DOCX处理失败: {str(e)}')


def process_epub(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    处理EPUB文件，转换为Markdown
    :param file_path: EPUB文件路径
    :param file_name: 文件名
    :return: (Markdown内容, 转换后的文件名)
    """
    if not HAS_EPUB:
        logger.error(f'ebooklib/BeautifulSoup 未安装，无法处理EPUB文件: {file_name}')
        raise ImportError('ebooklib 和 beautifulsoup4 未安装，请运行: pip install ebooklib beautifulsoup4')
    
    try:
        book = epub.read_epub(file_path)
        content_parts = []
        
        # 获取书籍标题
        title = None
        for item in book.get_items():
            if item.get_name() == 'title':
                title = item.get_content().decode('utf-8')
                break
        
        if title:
            content_parts.append(f"# {title}\n\n")
        
        # 按spine顺序提取章节
        spine_items = list(book.spine)
        
        for item_id, _ in spine_items:
            item = book.get_item_by_id(item_id)
            if item is None:
                continue
            
            # 只处理HTML/XHTML内容
            content_type = item.get_type()
            if content_type not in [ebooklib.ITEM_HTML, ebooklib.ITEM_XHTML, ebooklib.ITEM_DOCUMENT]:
                continue
            
            # 解析HTML内容
            html_content = item.get_content().decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取标题
            title_tag = soup.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'title'])
            if title_tag:
                title_text = title_tag.get_text().strip()
                if title_text and title_text != title:
                    level = int(title_tag.name[1]) if title_tag.name.startswith('h') else 2
                    content_parts.append(f"{'#' * level} {title_text}\n\n")
            
            # 提取正文
            body = soup.find('body') or soup
            text_content = body.get_text(separator='\n', strip=True)
            
            if text_content:
                # 清理文本
                text_content = re.sub(r'\n{3,}', '\n\n', text_content)
                content_parts.append(f"{text_content}\n\n")
        
        content = "\n".join(content_parts)
        
        # 清理内容
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # 生成Markdown文件名
        md_file_name = file_name.replace('.epub', '.md')
        
        logger.info(f'EPUB处理完成: {file_name}, 提取内容长度: {len(content)} 字符')
        return content.strip(), md_file_name
    except Exception as e:
        logger.error(f'EPUB处理失败: {file_name}, 错误: {str(e)}', exc_info=True)
        raise Exception(f'EPUB处理失败: {str(e)}')


def process_text(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    处理文本文件（Markdown、TXT）
    :param file_path: 文件路径
    :param file_name: 文件名
    :return: (文件内容, 文件名)
    """
    try:
        logger.debug(f'开始处理文本文件: {file_path}')
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            logger.error(f'无法解码文件内容: {file_name}, 尝试的编码: {encodings}')
            raise ValueError('无法解码文件内容')
        
        logger.debug(f'文本文件读取成功: {file_name}, 编码: {used_encoding}, 内容长度: {len(content)} 字符')
        
        # 如果是TXT文件，转换为Markdown
        if file_name.endswith('.txt'):
            md_file_name = file_name.replace('.txt', '.md')
        else:
            md_file_name = file_name
        
        logger.info(f'文本文件处理完成: {file_name}, Markdown文件名: {md_file_name}')
        return content.strip(), md_file_name
    except Exception as e:
        logger.error(f'文本文件处理失败: {file_name}, 错误: {str(e)}', exc_info=True)
        raise Exception(f'文本文件处理失败: {str(e)}')


def save_processed_file(original_path: str, content: str, md_file_name: str) -> str:
    """
    保存处理后的Markdown文件
    :param original_path: 原始文件路径
    :param content: Markdown内容
    :param md_file_name: Markdown文件名
    :return: 保存后的文件路径
    """
    file_dir = Path(original_path).parent
    md_file_path = file_dir / md_file_name
    
    logger.debug(f'保存Markdown文件: {md_file_path}, 内容长度: {len(content)} 字符')
    with open(md_file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    logger.info(f'Markdown文件保存成功: {md_file_path}')
    return str(md_file_path)

